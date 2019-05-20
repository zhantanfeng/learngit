# -*- coding: UTF-8 -*-
from nameko.rpc import rpc
import pymysql
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_UNDERLINE, WD_LINE_SPACING
from docx.oxml.ns import qn

class Compute(object):
    name = "test"

    @rpc
    def test(self):
        return "hello"

    @rpc
    def compute(self):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select TEACHER_ID from es_honor where HONOR like '%院士%' "
        cursor.execute(sql)
        teacher_id = cursor.fetchall()
        return teacher_id

    @rpc
    def compute1(self,id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,charset='utf8')
        cursor = db.cursor()
        sql = "select NAME from es_teacher where ID = %s "
        cursor.execute(sql,(id))
        teacher_name = cursor.fetchone()
        return teacher_name

    @rpc
    def compute2(self):
        teacher_name = []
        teacher_id = self.compute()
        for i in teacher_id:
            teacher_name.append(self.compute1(i))
        return teacher_name

    @rpc
    def get_institutionId(self,schoolName,institutionName):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306, charset='utf8')
        cursor = db.cursor()
        sql = "select ID from es_institution where SCHOOL_NAME = %s and NAME = %s "
        cursor.execute(sql, (schoolName,institutionName))
        institution_id = cursor.fetchall()
        return institution_id

    @rpc
    def get_schoolId(self,schoolName):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select SCHOOL_ID from es_institution where SCHOOL_NAME = %s"
        cursor.execute(sql, (schoolName))
        school_id = cursor.fetchone()
        return school_id

    @rpc
    def get_teacher_name_and_insId(self,school_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME,INSTITUTION_ID,HOMEPAGE from es_teacher where SCHOOL_ID = %s and ACADEMICIAN > 1"
        cursor.execute(sql, (school_id))
        teacher = cursor.fetchall()
        return teacher

    @rpc
    def get_institution_name(self,institution_ID):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME from es_institution where ID = %s "
        cursor.execute(sql, (institution_ID))
        institution_id = cursor.fetchone()
        return institution_id

    @rpc
    def get_academicianName(self,institution_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306, charset='utf8')
        cursor = db.cursor()
        sql = "select NAME, HOMEPAGE from es_teacher where INSTITUTION_ID = %s and ACADEMICIAN > 0 "
        cursor.execute(sql, (institution_id))
        teacher_name = cursor.fetchall()
        return teacher_name

    @rpc
    def get_institutionNamebyschoolName(self,schoolname):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME from es_institution where SCHOOL_NAME = %s "
        cursor.execute(sql, (schoolname))
        institution_name = cursor.fetchall()
        return institution_name

class document(object):
    name = "document"

    # 根据学校名，学院名获取学院ID
    @rpc
    def get_institutionId(self, schoolName, institutionName):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID from es_institution where SCHOOL_NAME = %s and NAME = %s "
        cursor.execute(sql, (schoolName, institutionName))
        institution_id = cursor.fetchone()
        institution_id = institution_id[0]
        return institution_id

    # 根据学院id获取学院中所有老师的ID,姓名，是否院士，是否杰出青年，是否长江学者
    @rpc
    def get_teacher_info(self, institutionId):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID,NAME,ACADEMICIAN,OUTYOUTH,CHANGJIANG from es_teacher where INSTITUTION_ID = %s"
        cursor.execute(sql, (institutionId))
        teacherInfo = cursor.fetchall()
        return teacherInfo

    # 根据学校名，学院名获取重点实验室名
    @rpc
    def get_lab(self, school_name, institution_name):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT name FROM main_lab where org= %s and institution = %s "
        cursor.execute(sql, (school_name, institution_name))
        lab = cursor.fetchall()
        lab_name = []
        for i in lab:
            b = i[0].index("（")
            lab_name.append(i[0][0:b])
        return lab_name

    # 获取领头人领域
    @rpc
    def get_fields(self, institution_id, teacher_name):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT FIELDS FROM es_teacher where INSTITUTION_ID = %s and NAME = %s "
        cursor.execute(sql, (institution_id, teacher_name))
        fields = cursor.fetchone()
        return fields

    # 根据学院id获取重点学科代码，评价
    @rpc
    def get_maindis(self, institution_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = '''SELECT DISCIPLINE_CODE,EVALUATION FROM es_relation_in_dis where INSTITUTION_ID = %s and (EVALUATION = 'A+' or EVALUATION = 'A')'''
        cursor.execute(sql, (institution_id))
        maindis = cursor.fetchall()
        return maindis

    # 根据老师id获取合著老师姓名，合著数量
    @rpc
    def get_relation(self, teacher_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select teacher2_name,paper_num from teacher_teacher where teacher1_id = %s"
        cursor.execute(sql, (teacher_id))
        teacher_list = cursor.fetchall()
        return teacher_list

    # 根据学科代码获取学科名
    @rpc
    def get_discipline(self, discipline_code):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT NAME FROM es_discipline where CODE = %s"
        cursor.execute(sql, (discipline_code))
        discipline_name = cursor.fetchone()
        discipline_name = discipline_name[0]
        return discipline_name

    # 根据老师名，学院id,获取老师ID
    @rpc
    def get_teacher_id(self, name, institution):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT ID FROM es_teacher where NAME = %s and INSTITUTION_ID = %s"
        cursor.execute(sql, (name, institution))
        teacher_id = cursor.fetchone()
        return teacher_id

    # 根据学校名获取带头人姓名，项目名，项目年份
    @rpc
    def get_project(self, org):
        db = pymysql.connect(host='47.104.236.183', db='eds_spider', user='root', password='SLX..eds123', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT PERSON,PROJECT_NAME,FUNDS,YEAR FROM eval_project where ORG = %s and FUNDS is not NULL"
        cursor.execute(sql, (org))
        project = cursor.fetchall()
        project_list = []
        for i in project:
            project_list.append(i)
        return project_list

    # 根据作者id,年份获取论文所有作者，论文名
    @rpc
    def get_paper_info_1(self, author_id, year):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT author,name from eds_paper_clean where author_id = %s and year = %s"
        cursor.execute(sql, (author_id, year))
        paper = cursor.fetchall()
        paperlist = []
        for i in paper:
            paperlist.append(i[0])
        return paperlist

    # 根据作者id,获取论文所有作者，论文名
    @rpc
    def get_paper_info_2(self, author_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT author,name from eds_paper_clean where author_id = %s"
        cursor.execute(sql, (author_id))
        paper = cursor.fetchall()
        # paperlist = []
        # for i in paper:
        #     paperlist.append(i[0])
        return paper

    # 根据老师ID，年份获取荣誉
    @rpc
    def get_honor_1(self, teacher_id, year):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select HONOR from es_honor where TEACHER_ID = %s and year = %s"
        cursor.execute(sql, (teacher_id, year))
        honor = cursor.fetchall()
        return honor

    # 根据老师ID，年份获取荣誉
    @rpc
    def get_honor_2(self, teacher_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select HONOR from es_honor where TEACHER_ID = %s"
        cursor.execute(sql, (teacher_id))
        honor = cursor.fetchall()
        return honor

    @rpc
    # 根据老师名获取老师头衔：是否院士，是否长江，是否杰青
    def get_title(self, teacherName, institution_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME,ACADEMICIAN,OUTYOUTH,CHANGJIANG from es_teacher where NAME = %s and INSTITUTION_ID=%s"
        cursor.execute(sql, (teacherName, institution_id))
        title = cursor.fetchone()
        return title

    @rpc
    # 根据老师名获取老师专利
    def get_invention(self, teacherName):
        db = pymysql.connect(host='47.104.236.183', db='lw_temp', user='root', password='SLX..eds123', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select title,date1 from cnki_zhuanli where author_list like '%%%s%%'" %(teacherName)
        cursor.execute(sql)
        invention = cursor.fetchall()
        return invention

    # 创建文档
    @rpc
    def createdocument(self,institution_info, team):
        document = Document()
        document.styles['Normal'].font.name = u'微软雅黑'
        document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        # 创建段落
        p = document.add_paragraph("")
        # 设置段落左右居中
        #p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置段落的段前间距
        p.paragraph_format.space_before = Pt(5)
        # 设置段落得断后间距
        p.paragraph_format.space_after = Pt(5)
        # 设置行间距
        p.paragraph_format.line_spacing = Pt(8)
        # 设置段落间距的格式为最小值
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        run = p.add_run(
            institution_info["school_name"] + "科研简报" + institution_info["institution_name"] + institution_info["date"])
        # 设置字体
        run.font.name = u'宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        # 颜色设置，这里是用RGB颜色
        run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置字体大小
        run.font.size = Pt(21)
        # 字体是否加粗
        run.bold = True
        # 无下划线
        run.underline = WD_UNDERLINE.NONE

        # 创建表格
        table = document.add_table(rows=1, cols=1)
        # 表格风格
        table.style = "Table Grid"

        for row in table.rows:
            # 设置每行表格高度
            row.height = Pt(30)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        # 在表格中写入文字
        run = table.cell(0, 0).paragraphs[0].add_run("一、院系概况")
        # 设置表格中字体

        # 字体大小
        run.font.size = Pt(15)
        # 字体大小
        run.font.color.rgb = RGBColor(91, 155, 213)
        # 是否加粗
        run.bold = True

        table = document.add_table(rows=1, cols=2)
        # 表格风格
        table.style = "Table Grid"

        for row in table.rows:
            # 设置每行表格高度
            row.height = Pt(30)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        run = table.cell(0, 0).paragraphs[0].add_run("国家重点学科")
        table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
        table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
        table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(237, 125, 49)
        run = table.cell(0, 1).paragraphs[0].add_run("评价")
        table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(2)
        table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
        table.cell(0, 1).paragraphs[0].paragraph_format.line_spacing = Pt(30)



        run.font.size = Pt(14)
        run.bold = True
        run.font.color.rgb = RGBColor(237, 125, 49)
        # 创建表格
        table = document.add_table(rows=len(institution_info["maindis"]), cols=2)
        # 表格风格
        table.style = "Table Grid"

        for row in table.rows:
            # 设置每行表格高度
            row.height = Pt(30)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST

        key = []
        value = []
        for i in institution_info["maindis"].keys():
            key.append(i)
        for i in institution_info["maindis"].values():
            value.append(i)
        count = len(institution_info["maindis"])

        for i in range(0, count):
            run = table.cell(i, 0).paragraphs[0].add_run(key[i])
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            table.cell(i, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(i, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(i, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            run = table.cell(i, 1).paragraphs[0].add_run(value[i])
            run.font.name = u'微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            table.cell(i, 1).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(i, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(i, 1).paragraphs[0].paragraph_format.line_spacing = Pt(30)


        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        for row in table.rows:
            row.height = Pt(30)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        run = table.cell(0, 0).paragraphs[0].add_run("科研平台")

        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(237, 125, 49)
        run.bold = True
        table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
        table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
        table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)


        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        for row in table.rows:
            row.height = Pt(30)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        run = table.cell(0, 0).paragraphs[0].add_run(institution_info["mainlab"])

        run.font.size = Pt(12)
        table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
        table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
        table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

        # 团队介绍
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        for row in table.rows:
            row.height = Pt(30)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        run = table.cell(0, 0).paragraphs[0].add_run("二、科研团队")
        table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
        table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
        table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(91, 155, 213)
        run.bold = True


        for i in team:
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run("项目成员")

            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(237, 125, 49)
            run.bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            # 领军人物
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.cell(0, 0)
            table.cell(0, 1).merge(table.cell(0, 2)).merge(table.cell(0, 3)).merge(table.cell(0, 4))

            run = table.cell(0, 0).paragraphs[0].add_run("领军人物")
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)
            run.font.size = Pt(12)
            run.bold = True
            run = table.cell(0, 2).paragraphs[0].add_run(i['head_name'])

            run.font.size = Pt(12)
            table.cell(0, 2).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 2).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 2).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            if len(i["changjiang_list"]) > 0:
                table = document.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                table.cell(0, 0)
                table.cell(0, 1).merge(table.cell(0, 2)).merge(table.cell(0, 3)).merge(table.cell(0, 4))
                run = table.cell(0, 0).paragraphs[0].add_run("院士")
                run.font.size = Pt(12)
                run.bold = True
                table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
                table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
                table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

                run = table.cell(0, 1).paragraphs[0].add_run(j + " " for j in i['academician_list'])

                run.font.size = Pt(12)
                table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(2)
                table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
                table.cell(0, 1).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            if len(i["changjiang_list"]) > 0:
                table = document.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                table.cell(0, 0)
                table.cell(0, 1).merge(table.cell(0, 2)).merge(table.cell(0, 3)).merge(table.cell(0, 4))

                run = table.cell(0, 0).paragraphs[0].add_run("长江学者")
                run.font.size = Pt(12)
                run.bold = True
                table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
                table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
                table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)


                run = table.cell(0, 1).paragraphs[0].add_run(j + " " for j in i["changjiang_list"])

                run.font.size = Pt(12)
                table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(2)
                table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
                table.cell(0, 1).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            if len(i['outyouth_list']) > 0:
                table = document.add_table(rows=1, cols=5)
                table.style = "Table Grid"
                table.cell(0, 0)
                table.cell(0, 1).merge(table.cell(0, 2)).merge(table.cell(0, 3)).merge(table.cell(0, 4))
                run = table.cell(0, 0).paragraphs[0].add_run("杰出青年")
                table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
                table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
                table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

                run.font.size = Pt(12)
                run.bold = True

                run = table.cell(0, 1).paragraphs[0].add_run(j + " " for j in i["outyouth_list"])

                run.font.size = Pt(12)
                table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(2)
                table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
                table.cell(0, 1).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            # 其他成员
            table = document.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            table.cell(0, 0)
            table.cell(0, 1).merge(table.cell(0, 2)).merge(table.cell(0, 3)).merge(table.cell(0, 4))
            run = table.cell(0, 0).paragraphs[0].add_run("其他成员")
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)
            run.font.size = Pt(12)
            run.bold = True


            run = table.cell(0, 1).paragraphs[0].add_run(j + " " for j in i["other_list"])

            run.font.size = Pt(12)
            table.cell(0, 1).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 1).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 1).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            # 团队研究方向
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run("团队研究方向")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(237, 125, 49)
            run.bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run(i["team_direction"])

            run.font.size = Pt(12)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            # 团队成果
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run("团队成果")

            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(237, 125, 49)
            run.bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            # 论文成果
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run("论文成果")

            run.font.size = Pt(12)
            run.bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            run = table.cell(0, 0).paragraphs[0].add_run("《" + j + "》" for j in i["paper"])

            run.font.size = Pt(12)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            # 专利成果
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run("专利成果")

            run.font.size = Pt(12)
            run.bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            run = table.cell(0, 0).paragraphs[0].add_run("《" + j + "》" for j in i["invention"])

            run.font.size = Pt(12)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)


            # 获奖成果
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            run = table.cell(0, 0).paragraphs[0].add_run("获奖成果")

            run.font.size = Pt(12)
            run.bold = True
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            run = table.cell(0, 0).paragraphs[0].add_run(j for j in i["award"])

            run.font.size = Pt(12)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_before = Pt(2)
            table.cell(0, 0).paragraphs[0].paragraph_format.space_after = Pt(6)
            table.cell(0, 0).paragraphs[0].paragraph_format.line_spacing = Pt(30)

            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            for row in table.rows:
                row.height = Pt(30)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


        # 保存文档
        document.save("./static/docx/"+institution_info['school_name']+"科研简报"+institution_info['institution_name']+institution_info['date']+".docx")


class team(object):
    name = "team"

    @rpc
    #根据学校名和学院名获取学院id
    def get_institutionId(self, schoolName, institutionName):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID,SCHOOL_ID from es_institution where SCHOOL_NAME = %s and NAME = %s "
        cursor.execute(sql, (schoolName, institutionName))
        institution = cursor.fetchone()
        return institution

    @rpc
    #根据学校名获取学院名获取这个学院所有老师
    def get_teacher(self,school_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME from es_teacher where SCHOOL_ID = %s"
        cursor.execute(sql,(school_id))
        teacher_list = cursor.fetchall()
        return teacher_list

    @rpc
    #根据老师名，学校名，学院名获取老师的id
    def get_teacher_id(self,teacher_name,school_id,institution_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID from es_teacher where NAME = %s and SCHOOL_ID = %s and INSTITUTION_ID = %s"
        cursor.execute(sql,(teacher_name,school_id,institution_id))
        teacher_id = cursor.fetchone()
        return teacher_id

    @rpc
    #根据老师id,获取所有论文的作者
    def get_member(self,author_id):
        db = pymysql.connect(host='47.104.236.183', db='eds_base', user='root', password='SLX..eds123', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql  = "select author from es_paper where author_id = %s"
        cursor.execute(sql,(author_id))
        member = cursor.fetchall()
        return member

    @rpc
    #根据老师名获取老师头衔：是否院士，是否长江，是否杰青
    def get_title(self, teacherName ,institution_id):
        db = pymysql.connect(host='47.106.83.33',db='eds_base',user='root',password='111111',port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME,ACADEMICIAN,OUTYOUTH,CHANGJIANG from es_teacher where NAME = %s and INSTITUTION_ID=%s"
        cursor.execute(sql, (teacherName, institution_id))
        title = cursor.fetchone()
        return title



class title_search(object):
    name = "title_search"
    @rpc
    # 根据学校名和学院名获取学院id
    def get_institutionId(self, schoolName, institutionName):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID from es_institution where SCHOOL_NAME = %s and NAME = %s "
        cursor.execute(sql, (schoolName, institutionName))
        institution_id = cursor.fetchone()
        return institution_id

    @rpc
    #根据老师名获取老师头衔：是否院士，是否长江，是否杰青
    def get_title(self, teacherName ,institution_id):
        db = pymysql.connect(host='47.106.83.33',db='eds_base',user='root',password='111111',port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select NAME,ACADEMICIAN,OUTYOUTH,CHANGJIANG from es_teacher where NAME = %s and INSTITUTION_ID=%s"
        cursor.execute(sql, (teacherName, institution_id))
        title = cursor.fetchone()
        return title

class paper_search(object):
    name = "paper_search"

    @rpc
    # 根据学校名和学院名获取学院id
    def get_institutionId(self, schoolName, institutionName):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID from es_institution where SCHOOL_NAME = %s and NAME = %s "
        cursor.execute(sql, (schoolName, institutionName))
        institution_id = cursor.fetchone()
        return institution_id

    @rpc
    #根据老师名和学院id获取老师的id
    def get_teacherid(self,teacher_name,institution_id):
        db = pymysql.connect(host='47.106.83.33', db='eds_base', user='root', password='111111', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "select ID from es_teacher where NAME = %s and INSTITUTION_ID = %s"
        cursor.execute(sql,(teacher_name,institution_id))
        teacher_id = cursor.fetchone()
        return teacher_id

   # 根据作者id,获取所有论文名和年份
    @rpc
    def get_paper(self, author_id):
        db = pymysql.connect(host='47.104.236.183', db='eds_base', user='root', password='SLX..eds123', port=3306,
                             charset='utf8')
        cursor = db.cursor()
        sql = "SELECT name,year from es_paper where author_id = %s"
        cursor.execute(sql, (author_id))
        paper = cursor.fetchall()
        return paper
    
    






